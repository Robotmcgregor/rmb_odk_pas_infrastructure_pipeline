#!/usr/bin/env python

"""
Copyright 2021 Robert McGregor

Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated
documentation files (the "Software"), to deal in the Software without restriction, including without limitation the
rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software,
and to permit persons to whom the Software is furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all copies or substantial portions of the
Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE
WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NON INFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR
COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT,
TORT OR OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""

# Import modules
from __future__ import print_function, division
import warnings
from glob import glob
import pandas as pd
import os
import urllib
from urllib import request
from docx import Document
from docx.shared import Cm
import geopandas as gpd
import matplotlib.pyplot as plt

warnings.filterwarnings("ignore")


def user_id_fn(user_df):
    """ Extract the user id stripping of adm when on the remote desktop.

    :param remote_desktop: string object containing the computer system being used.
    :return final_user: string object containing the NTG user id. """

    # extract user name
    home_dir = os.path.expanduser("~")
    _, user = home_dir.rsplit('\\', 1)

    single_user_df = user_df[user_df['user_id'] == user]
    name = single_user_df.name.iloc[0]
    email = single_user_df.email.iloc[0]
    phone = single_user_df.phone.iloc[0]

    final_user = user[3:]

    return name, email, phone


def specimen_name_fn(uid_df):
    import numpy as np

    sample1_label = str(uid_df['sample1_label'].iloc[0])
    sample2_label = str(uid_df['sample2_label'].iloc[0])
    sample3_label = str(uid_df['sample3_label'].iloc[0])

    sample_label_list = []
    if sample1_label != 'Nan':
        sample_label_list.append(sample1_label)

    if sample2_label != 'Nan':
        sample_label_list.append(sample2_label)

    if sample3_label != 'Nan':
        sample_label_list.append(sample3_label)

    sample_names = ', '.join(sample_label_list)
    final_sample_name = sample_names[0:]

    return final_sample_name, sample_label_list


def photos_download_fn(uid_df, document, export_dir, sample_label_list):
    for i in range(3):

        sample_photo = str(uid_df['sample' + str(i + 1) + '_photo'].iloc[0])

        if sample_photo != 'nan':
            sample_name = str(sample_label_list[i]).replace(' ', '_')

            output_str = export_dir + '\\sample' + str(i) + '_' + sample_name + '_photo.jpg'
            request.urlretrieve(sample_photo, output_str)
            document.add_picture(output_str, width=Cm(6))
            document.add_paragraph('Sample1: , ' + str(sample_label_list[i]))

        i += 1

    return document


def photo_download_and_insertion_range_loop_fn(column_name, photo_file_name, text_field, document, uid_df, export_dir,
                                               uid, prop_code, date_label):
    for n in range(3):

        if str(uid_df[column_name + str(n + 1)].iloc[0]) != 'nan':
            output_str = export_dir + '\\' + prop_code + '_' + date_label + '_' + photo_file_name + '_uid' + str(
                uid) + '_photo.jpg'
            request.urlretrieve(str(uid_df[column_name + str(n + 1)].iloc[0]), output_str)
            document.add_picture(output_str, width=Cm(6))
            paragraph = document.add_paragraph(text_field + (str(n + 1)))

    return document


def photo_download_and_insertion_range_loop2_fn(column_name, column_name_end, photo_file_name, text_field, document,
                                                uid_df, export_dir, uid, prop_code, date_label):
    for n in range(3):

        if str(uid_df[column_name + str(n + 1) + '_' + column_name_end].iloc[0]) != 'nan':
            if n == 0:
                paragraph = document.add_paragraph('The following samples were taken.')
            output_str = export_dir + '\\' + prop_code + '-' + date_label + '_' + photo_file_name + '_uid' + str(
                uid) + '_photo.jpg'
            request.urlretrieve(str(uid_df[column_name + str(n + 1) + '_' + column_name_end].iloc[0]), output_str)
            document.add_picture(output_str, width=Cm(6))
            paragraph = document.add_paragraph(text_field + (str(n + 1)))

    return document


def single_photo_download_and_insertion_fn(photo_type, export_dir, uid, uid_df, document, prop_code, date_label):
    if str(uid_df[photo_type + '_photo'].iloc[0]) == 'nan':
        pass
    elif str(uid_df[photo_type + '_photo'].iloc[0]) == 'Nan':
        pass
    else:
        output_str = export_dir + '\\' + prop_code + '_' + date_label + '_' + photo_type + '_uid' + str(
            uid) + '_photo.jpg'
        request.urlretrieve(str(uid_df[photo_type + '_photo'].iloc[0]), output_str)
        document.add_picture(output_str, width=Cm(6))

    return document


def single_heading_para_photo_fn(document, uid_df, uid, export_dir, item, prop_code, date_label):
    document.add_heading(item.title(), level=2)
    item_ = str(uid_df[item].iloc[0])
    document.add_paragraph(item + ': ' + str(item_))
    # call the single_photo_download_and_insertion_fn function to download and insert images into the document
    document = single_photo_download_and_insertion_fn(item, export_dir, uid, uid_df, document, prop_code, date_label)

    return document


def single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, item_list, photo_yn, prop_code, date_label):
    document.add_heading(item_list[0].title(), level=2)

    for item in item_list:

        variable = str(uid_df[item].iloc[0])

        if variable != 'Nan':
            variable_ = variable
        else:
            variable_ = 'Not recorded'

        document.add_paragraph(item.title() + ': ' + str(variable_))
        # call the single_photo_download_and_insertion_fn function to download and insert images into the document

        if photo_yn == 'Yes':
            document = single_photo_download_and_insertion_fn(item, export_dir, uid, uid_df, document, prop_code,
                                                              date_label)
        else:
            pass
    return document


def date_label_fn(date):
    year, month, day = date.split('-')
    if len(day) == 1:
        final_day = str(0) + str(day)
    else:
        final_day = str(day)

    if len(month) == 1:
        final_month = str(0) + str(month)
    else:
        final_month = str(month)

    date_label = str(year) + str(final_month + str(final_day))

    return date_label


def plot_point_on_property(pastoral_estate, prop_code, export_dir, uid_df, uid):
    # extract eastings and northing values for geometry feature
    '''row["easting"] = row["geometry"].x
    row["northing"] = row["geometry"].y'''

    prop_gdf = pastoral_estate[pastoral_estate["PROP_TAG"] == prop_code]

    if len(prop_gdf.index) > 0:

        # extract property name in title case.
        prop_name = (pastoral_estate.loc[pastoral_estate["PROP_TAG"] == prop_code, "PROPERTY"].item()).title()

        # create two subplots
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6), constrained_layout=True)
        fig.suptitle('Locality and location of unidentified species (' + prop_name + ')', fontsize=20)

        prop_bounds = prop_gdf.geometry.total_bounds
        prop_gdf.reset_index(drop=True, inplace=True)

        # set plot styles
        pastoral_estate.plot(ax=ax1, facecolor='gray')
        prop_gdf.plot(ax=ax1, facecolor='red')

        prop_bounds = prop_gdf.geometry.total_bounds

        # set plot extent based on the property +- 0.8 dd.
        xlim = ([prop_gdf.total_bounds[0] - 0.08, prop_gdf.total_bounds[2] + 0.08])
        ylim = ([prop_gdf.total_bounds[1] - 0.08, prop_gdf.total_bounds[3] + 0.08])
        ax2.set_xlim(xlim)
        ax2.set_ylim(ylim)
        # set plot styles
        pastoral_estate.plot(ax=ax2, facecolor='gray')

        prop_gdf.plot(ax=ax2, facecolor='red')

        uid_df.plot(ax=ax2, color='black', markersize=8)

        # set the spacing between subplots
        plt.subplot_tool()

        # export plots
        export_file_str = export_dir + "\\" + prop_code + '_location_uid' + str(uid) + ".jpg"
        plt.tight_layout()
        plt.savefig(export_file_str)

    else:
        export_file_str = 'nan'
        prop_name = 'Non_pastoral_property'
        pass

    return export_file_str, prop_name


def main_routine(dest52, dest53, export_dir, pastoral_estate, user_df):
    from docx import Document
    import numpy as np
    form_dir = export_dir + '\\request_id_forms'

    if not os.path.exists(form_dir):
        os.mkdir(form_dir)

    export_dir = export_dir + "\\photos\\"
    user_df = pd.read_csv(r'E:\DEPWS\code\rangeland_monitoring\rmb_mapping_pipeline\assets\contact_details.csv')

    i = 1
    for zone_df in [dest52, dest53]:
        zone_df['uid'] = zone_df.index + 1
        # unidentified_df = gdf.to_crs(epsg=4283)

        # extract eastings and northing values for geometry feature
        zone_df["longitude"] = zone_df["geometry"].x
        zone_df["latitude"] = zone_df["geometry"].y

        for uid in zone_df.uid.unique():
            # for index, row in unidentified_df.iterrows():
            uid_df = zone_df[zone_df['uid'] == uid]

            list(uid_df)
            date = uid_df['date_rec'].iloc[0]
            date_label = date_label_fn(date)

            document = Document(r'E:\DEPWS\code\rangeland_monitoring\rmb_mapping_pipeline\assets\species_request.docx')

            habit_ = uid_df['habit'].iloc[0]
            prop_code = uid_df['prop_code'].iloc[0]
            document.add_heading('Species identification request', level=1)
            paragraph = document.add_paragraph('Division/Branch: Rangeland Monitoring Branch.')

            # call the user id function to extract the user id
            name, email, phone = user_id_fn(user_df)

            document.add_paragraph('Officers name: ' + str(name))
            document.add_paragraph('Phone: ' + str(phone))
            document.add_paragraph('Email: ' + str(email))

            document.add_heading('Location', level=2)
            # call the
            sample_names, sample_label_list = specimen_name_fn(uid_df)
            document.add_paragraph('Specimen reference name: ' + str(sample_names))
            longitude = zone_df['longitude'].loc[0]
            document.add_paragraph('Longitude (GDA94): ' + str(longitude))
            latitude = zone_df['latitude'].loc[0]
            document.add_paragraph('Latitude (GDA94): ' + str(latitude))

            output_file_str, prop_name = plot_point_on_property(pastoral_estate, prop_code, export_dir, uid_df, uid)
            if output_file_str != 'nan':
                document.add_picture(output_file_str, width=Cm(18))
            # --------------------------------------------- Habit ------------------------------------------------------
            document.add_heading('Habit', level=2)

            abundance = float(uid_df['abundance'].iloc[0])
            if float(abundance) > 1:
                document.add_paragraph('Specimen abundance: ' + str(abundance))
            else:
                document.add_paragraph('Specimen abundance: ' + str('Not recorded'))

            density = float(uid_df['density'].iloc[0])
            if density > 1:
                document.add_paragraph('Specimen density: ' + str(density))
            else:
                document.add_paragraph('Specimen density: ' + str('Not recorded'))

            # call the photo_download_and_insertion_range_loop_fn function to loop through the three possible species photos.
            document = photo_download_and_insertion_range_loop_fn(
                'photo', 'habit', 'Figure x: Species in situ', document, uid_df, export_dir, uid, prop_code, date_label)

            # call the photo_download_and_insertion_range_loop2_fn function to loop through the three possible sample photos.
            photo_download_and_insertion_range_loop2_fn(
                'sample', 'photo', 'Sample ', 'Figure x: Sample photograph', document, uid_df, export_dir, uid,
                prop_code, date_label)

            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir,
                                                         ['habit', 'height', 'annual_perennial'], 'No', prop_code,
                                                         date_label)
            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['roots'], 'Yes', prop_code,
                                                         date_label)
            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['cork', 'cork_colour'],
                                                         'Yes', prop_code, date_label)
            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['leaf'], 'Yes', prop_code,
                                                         date_label)
            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['flower_colour'], 'Yes',
                                                         prop_code, date_label)
            document = single_photo_download_and_insertion_fn('flower_cluster', export_dir, uid, uid_df, document,
                                                              prop_code, date_label)

            # ------------------------------------------------ grass ---------------------------------------------------

            if habit_ == 'Grass':
                document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['grass_head'], 'Yes',
                                                             prop_code, date_label)
                document = single_photo_download_and_insertion_fn('grass_flower', export_dir, uid, uid_df, document,
                                                                  prop_code, date_label)
                document = single_photo_download_and_insertion_fn('grass_spikelet', export_dir, uid, uid_df, document,
                                                                  prop_code, date_label)
                document = single_photo_download_and_insertion_fn('grass_awn', export_dir, uid, uid_df, document,
                                                                  prop_code, date_label)

            # --------------------------------------------- fruit/seed -------------------------------------------------

            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['fruit'], 'Yes', prop_code,
                                                         date_label)
            document = single_photo_download_and_insertion_fn('seed', export_dir, uid, uid_df, document, prop_code,
                                                              date_label)

            # ------------------------------------------------ Odour ---------------------------------------------------
            document.add_heading('Odour', level=2)
            odour = str(uid_df['odour_yn'].iloc[0])

            if odour == 'Nan':
                document.add_paragraph('Odour:  ' + str('Not recorded'))

            else:
                document.add_paragraph('Odour:  ' + str('odour'))
                odour_desc = str(uid_df['odour'].iloc[0])
                if odour_desc != 'Nan':
                    document.add_paragraph('The smell is described as:  ' + str('odour_desc'))

            # ------------------------------------------------- prickles -----------------------------------------------

            document.add_heading('Prickles/Hairs', level=2)
            prickles = str(uid_df['prickles'].iloc[0])

            if prickles == 'Nan':
                document.add_paragraph('Prickles/Hairs:  ' + str('Not recorded'))

            elif prickles == 'Hairs present':
                document.add_paragraph('Prickles/Hairs:  ' + str(prickles))
                prick_location = str(uid_df['prickles_location'].iloc[0])

                if prick_location != 'Nan':
                    document.add_paragraph('Prickles/Hairs location:  ' + str(prick_location))

                else:
                    document.add_paragraph('Prickles/Hairs location:  ' + str('Not recorded'))

            else:
                document.add_paragraph('Prickles/Hairs:  ' + str(prickles))
                prick_location = str(uid_df['prickles_location'].iloc[0])
                if prick_location != 'Nan':
                    document.add_paragraph('Prickles/Hairs location:  ' + str(prick_location))

                else:
                    document.add_paragraph('Prickles/Hairs location:  ' + str('Not recorded'))

            # --------------------------------------------- latex -----------------------------------------------------

            document = single_heading_para_photo_list_fn(document, uid_df, uid, export_dir, ['latex'], 'Yes', prop_code,
                                                         date_label)

            if i == 1:

                zone_label = 'zone52'
            else:

                zone_label = 'zone53'

            document.save(form_dir + '\\Unidentified_species_' +
                          str(uid) + '_' + prop_name.replace(' ', '_') + '_' + str(zone_label) + '.docx')

        i += 1


if __name__ == '__main__':
    main_routine()
